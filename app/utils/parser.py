import io
import re
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document

# -------------------------------
# OCR helper
# -------------------------------
def ocr_page(page):
    try:
        pil_img = page.to_image(resolution=300).original
        return pytesseract.image_to_string(pil_img)
    except Exception:
        return ""

# -------------------------------
# PDF extractor (text + OCR fallback)
# -------------------------------
def extract_text_from_pdf(path_or_fileobj):
    text_parts = []

    with pdfplumber.open(path_or_fileobj) as pdf:
        for page in pdf.pages:
            # Try normal extraction
            text = page.extract_text()

            if text and text.strip():
                text_parts.append(text)
            else:
                # Fallback to OCR
                ocr_text = ocr_page(page)
                if ocr_text.strip():
                    text_parts.append(ocr_text)

    return "\n".join(text_parts)


# -------------------------------
# DOCX extractor
# -------------------------------
def extract_text_from_docx(path_or_fileobj):
    if hasattr(path_or_fileobj, "read"):
        doc = Document(io.BytesIO(path_or_fileobj.read()))
    else:
        doc = Document(path_or_fileobj)

    texts = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(texts)

# -------------------------------
# Clean text
# -------------------------------
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

# -------------------------------
# EMAIL + NAME extraction
# -------------------------------
def extract_email(text):
    m = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return m.group(0) if m else None

def infer_name_from_filename(filename):
    import os
    base = os.path.splitext(os.path.basename(filename))[0]
    base = base.replace("_", " ").replace("-", " ").strip()
    return base.title() if base else "Unknown"

def extract_years_experience(text):
    m = re.search(r"(\d{1,2})\+?\s+(?:years|yrs)\b", text.lower())
    if m:
        return int(m.group(1))
    return None

# -------------------------------
# Main function
# -------------------------------
def parse_resume(uploaded_file):
    filename = getattr(uploaded_file, "name", "uploaded_resume")

    # Choose extractor
    if filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(uploaded_file)
    elif filename.lower().endswith((".doc", ".docx")):
        text = extract_text_from_docx(uploaded_file)
    else:
        try:
            uploaded_file.seek(0)
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        except:
            text = ""

    text = clean_text(text)

    return {
        "filename": filename,
        "text": text,
        "email": extract_email(text),
        "name": infer_name_from_filename(filename),
        "years_experience": extract_years_experience(text),
    }
